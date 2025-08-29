import io
import datetime as dt
from typing import List, Dict

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==========================
# Estilos e utilidades DOCX
# ==========================

def _set_default_styles(doc: Document):
    styles = doc.styles
    # Fonte padrão
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        if style_name in styles:
            style = styles[style_name]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)
            # Para compatibilidade com MS Word
            try:
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            except Exception:
                pass

    if "Title" in styles:
        styles["Title"].font.size = Pt(20)
        styles["Title"].font.bold = True

    if "Heading 1" in styles:
        styles["Heading 1"].font.size = Pt(14)
        styles["Heading 1"].font.bold = True

    if "Heading 2" in styles:
        styles["Heading 2"].font.size = Pt(12)
        styles["Heading 2"].font.bold = True


def _add_header_block(doc: Document, org_name: str, logo_bytes: bytes | None):
    if logo_bytes:
        # Tentar inserir logo no cabeçalho superior
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        try:
            run.add_picture(io.BytesIO(logo_bytes), width=Inches(1.2))
        except Exception:
            pass
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if org_name.strip():
        p = doc.add_paragraph()
        run = p.add_run(org_name.strip())
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _human_date(date_obj: dt.date, hour_str: str | None = None):
    meses = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    d = date_obj.day
    m = meses[date_obj.month-1]
    y = date_obj.year
    if hour_str:
        return f"{d} de {m} de {y}, às {hour_str}"
    return f"{d} de {m} de {y}"


def _add_title(doc: Document, title: str):
    p = doc.add_paragraph(style="Title")
    p_run = p.add_run(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_heading(doc: Document, text: str):
    h = doc.add_paragraph(text, style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bullets(doc: Document, items: List[str], numbered: bool = False):
    """
    Cria lista com marcadores (padrão) ou numerada (numbered=True)
    usando estilos nativos do Word: 'List Bullet' e 'List Number'.
    Evita uso de APIs privadas do python-docx.
    """
    list_style = "List Number" if numbered else "List Bullet"
    for item in items:
        item = (item or "").strip()
        if not item:
            continue
        doc.add_paragraph(item, style=list_style)



def _add_actions_table(doc: Document, rows: List[Dict[str, str]]):
    if not rows:
        return
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Tarefa"
    hdr_cells[1].text = "Responsável"
    hdr_cells[2].text = "Prazo"

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = r.get("tarefa", "").strip()
        row_cells[1].text = r.get("responsavel", "").strip()
        row_cells[2].text = r.get("prazo", "").strip()

    # Ajuste simples de estilo
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]


# ==========================
# Geração da ata
# ==========================

from docx.oxml import OxmlElement
from docx.oxml.ns import qn as _qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def enable_line_numbering(doc: Document, count_by: int = 1, start: int = 1, restart: str = "continuous"):
    """
    Habilita numeração de linhas no documento (seção 0).
    restart: 'newPage' | 'newSection' | 'continuous'
    """
    sectPr = doc.sections[0]._sectPr
    ln = OxmlElement('w:lnNumType')
    ln.set(_qn('w:countBy'), str(count_by))
    ln.set(_qn('w:start'), str(start))
    ln.set(_qn('w:restart'), restart)
    sectPr.append(ln)


def add_justified_paragraph(doc: Document, text: str, first_line_indent_cm: float = 0.0):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent_cm > 0:
        p.paragraph_format.first_line_indent = Inches(first_line_indent_cm / 2.54)
    return p


def montar_narrativa(dados: Dict) -> list[str]:
    # Pauta como frase
    pauta = [i for i in dados.get('pauta', []) if i]
    pauta_frase = "; ".join(pauta) if pauta else "—"
    # Participantes
    participantes = [p for p in dados.get('participantes', []) if p]
    part_frase = ", ".join(participantes) if participantes else "—"
    # Encaminhamentos resumidos
    enc = []
    for r in dados.get('encaminhamentos', []):
        t = r.get('tarefa', '').strip(); resp = r.get('responsavel','').strip(); prazo = r.get('prazo','').strip()
        if any([t, resp, prazo]):
            bloco = " — ".join([x for x in [t, resp, prazo] if x])
            enc.append(bloco)
    enc_frase = "; ".join(enc) if enc else "—"

    p1 = (
        f"Ao {_human_date(dados['data'], dados.get('hora_inicio'))}, realizou-se a {dados.get('titulo','Reunião')} "
        f"no {dados.get('local','')}, tendo como pauta: {pauta_frase}."
    )

    p2 = (
        f"Estiveram presentes: {part_frase}. Presidiu: {dados.get('presidida_por','')}; "
        f"Secretariou: {dados.get('secretariada_por','')}."
    )

    p3 = f"Deliberações e registros: {dados.get('deliberacoes','').strip() or '—' }"
    p4 = f"Encaminhamentos: {enc_frase}."
    p5 = dados.get('encerramento','').strip()

    return [p for p in [p1, p2, p3, p4, p5] if p]



# ==========================
# UI – Streamlit
# ==========================
# ==========================

st.set_page_config(page_title="Gerador de Ata", layout="wide")

st.title("📝 Gerador de Ata de Reunião")

# Opções de formatação
col_fmt1, col_fmt2 = st.columns([1,1])
with col_fmt1:
    modo_narrativo = st.toggle("Formato narrativo (sem títulos de seção)", value=True, help="Gera parágrafos corridos.")
with col_fmt2:
    numerar_linhas = st.toggle("Numeração de linhas (Word)", value=True, help="Ativa numeração de linhas contínua no documento.")

with st.form("form_ata"):
    st.subheader("1) Cabeçalho")
    col1, col2 = st.columns([2,1])
    with col1:
        entidade = st.text_input("Entidade/Órgão (opcional)")
        titulo = st.text_input("Título da Ata", value="Ata de Reunião")
    with col2:
        logo_file = st.file_uploader("Logo (PNG/JPG opcional)", type=["png", "jpg", "jpeg"])

    st.subheader("2) Dados da Reunião")
    col3, col4, col5, col6 = st.columns([1,1,1,2])
    with col3:
        data = st.date_input("Data", value=dt.date.today())
    with col4:
        hora_inicio = st.text_input("Hora de início", value="09:00")
    with col5:
        hora_fim = st.text_input("Hora de término", value="11:00")
    with col6:
        local = st.text_input("Local", value="Sala de Reuniões")

    col7, col8 = st.columns(2)
    with col7:
        presidida_por = st.text_input("Presidida por")
    with col8:
        secretariada_por = st.text_input("Secretariada por")

    st.subheader("3) Participantes")
    participantes_str = st.text_area(
        "Um por linha",
        placeholder="Nome 1\nNome 2\nNome 3",
        height=120,
    )

    st.subheader("4) Pauta")
    pauta_str = st.text_area(
        "Itens da pauta (um por linha)",
        placeholder="1. Abertura\n2. Leitura da ata anterior\n3. Assuntos gerais",
        height=120,
    )

    st.subheader("5) Deliberações e Registros")
    deliberacoes = st.text_area(
        "Resumo das deliberações (texto livre)",
        placeholder="Descreva aqui os principais pontos debatidos e decisões tomadas…",
        height=160,
    )

    st.subheader("6) Encaminhamentos")
    st.caption("Preencha as linhas necessárias. Deixe em branco para ignorar.")
    enc_cols = st.columns([4,3,2])

    tarefas: List[Dict[str,str]] = []
    for i in range(1, 6):  # 5 linhas por padrão; pode aumentar futuramente
        with enc_cols[0]:
            tarefa = st.text_input(f"Tarefa {i}", key=f"tarefa_{i}")
        with enc_cols[1]:
            responsavel = st.text_input(f"Responsável {i}", key=f"resp_{i}")
        with enc_cols[2]:
            prazo = st.text_input(f"Prazo {i}", key=f"prazo_{i}")
        if any([tarefa.strip(), responsavel.strip(), prazo.strip()]):
            tarefas.append({"tarefa": tarefa, "responsavel": responsavel, "prazo": prazo})

    st.subheader("7) Encerramento")
    encerramento = st.text_area(
        "Texto final (opcional)",
        placeholder=(
            "Nada mais havendo a tratar, foi encerrada a reunião às "+
            (hora_fim or "__:")+
            ", e para constar eu, \"" + (secretariada_por or "________") + "\" lavrei a presente ata, que vai assinada."
        ),
        height=120,
    )

    st.subheader("8) Assinaturas")
    assinaturas_str = st.text_area(
        "Nomes para assinatura (um por linha, opcional)",
        placeholder="Presidente\nSecretário(a)",
        height=100,
    )

    gerar = st.form_submit_button("📥 Gerar .DOCX da Ata", use_container_width=True)

if 'last_doc' not in st.session_state:
    st.session_state['last_doc'] = None

if gerar:
    participantes = [x.strip() for x in (participantes_str or "").splitlines() if x.strip()]
    pauta = [x.strip() for x in (pauta_str or "").splitlines() if x.strip()]
    assinaturas = [x.strip() for x in (assinaturas_str or "").splitlines() if x.strip()]

    dados = {
        "entidade": entidade,
        "titulo": titulo,
        "data": data,
        "hora_inicio": hora_inicio,
        "hora_fim": hora_fim,
        "local": local,
        "presidida_por": presidida_por,
        "secretariada_por": secretariada_por,
        "participantes": participantes,
        "pauta": pauta,
        "deliberacoes": deliberacoes,
        "encaminhamentos": tarefas,
        "encerramento": encerramento,
        "assinaturas": assinaturas,
    }

    logo_bytes = logo_file.read() if logo_file else None

    # Monta o documento
    doc = Document()
    _set_default_styles(doc)
    if numerar_linhas:
        enable_line_numbering(doc, count_by=1, start=1, restart='continuous')

    # Cabeçalho + Título em CAIXA ALTA e centralizado (duas linhas se houver \" - \")
    _add_header_block(doc, dados.get("entidade", ""), logo_bytes)
    titulo_up = (dados.get("titulo") or "Ata de Reunião").upper()
    _add_title(doc, titulo_up)

    if modo_narrativo:
        # Corpo narrativo justificado com primeira linha recuada levemente
        for par in montar_narrativa(dados):
            add_justified_paragraph(doc, par, first_line_indent_cm=0.8)
        # Assinaturas
        _add_section_heading(doc, "ASSINATURAS")
        for a in dados.get("assinaturas", []):
            p = doc.add_paragraph("

")
            p = doc.add_paragraph(a)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Estrutura original com seções e listas (pauta numerada)
        _add_section_heading(doc, "1. IDENTIFICAÇÃO DA REUNIÃO")
        par = doc.add_paragraph()
        linha1 = f"Data: {_human_date(dados['data'], dados.get('hora_inicio'))} | Local: {dados.get('local','')}"
        par.add_run(linha1)
        par2 = doc.add_paragraph()
        linha2 = f"Presidida por: {dados.get('presidida_por','')} | Secretariada por: {dados.get('secretariada_por','')}"
        par2.add_run(linha2)

        _add_section_heading(doc, "2. PARTICIPANTES")
        _add_bullets(doc, participantes)

        _add_section_heading(doc, "3. PAUTA")
        _add_bullets(doc, pauta, numbered=True)

        _add_section_heading(doc, "4. DELIBERAÇÕES E REGISTROS")
        doc.add_paragraph(dados.get("deliberacoes", "") or "—")

        _add_section_heading(doc, "5. ENCAMINHAMENTOS")
        _add_actions_table(doc, dados.get("encaminhamentos", []))

        _add_section_heading(doc, "6. ENCERRAMENTO")
        if dados.get("encerramento"):
            doc.add_paragraph(dados["encerramento"])

        _add_section_heading(doc, "7. ASSINATURAS")
        for a in dados.get("assinaturas", []):
            p = doc.add_paragraph("

")
            p = doc.add_paragraph(a)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    st.session_state['last_doc'] = bio.getvalue()

# Área de download separada
st.markdown("---")
colA, colB = st.columns([1,2])
with colA:
    st.subheader("Exportar")
with colB:
    if st.session_state.get('last_doc'):
        filename = f"Ata_{dt.date.today().isoformat()}.docx"
        st.download_button(
            label="⬇️ Baixar arquivo .docx",
            data=st.session_state['last_doc'],
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.info("Preencha o formulário acima e clique em 'Gerar .DOCX da Ata'.")
